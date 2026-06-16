from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

root = Path.cwd()
img_dir = root / 'graficas_tfg' / 'data_fraction_gpu_barras'
out_path = root / 'informe_data_fraction_resumen.docx'

images = [
    ('Test MSE por cantidad de datos', img_dir / '01_barras_test_mse.png'),
    ('Test MSE por cantidad de datos (escala log)', img_dir / '02_barras_test_mse_log.png'),
    ('Violación física por cantidad de datos', img_dir / '03_barras_violacion_fisica.png'),
    ('Tiempo de inferencia por modelo', img_dir / '04_barras_tiempo_inferencia.png'),
    ('Mejor modelo por porcentaje de datos', img_dir / '05_barras_mejor_modelo_por_fraction.png'),
    ('Comparación con 100% de datos', img_dir / '06_barras_modelos_100_datos.png'),
    ('Comparación con 1% de datos', img_dir / '07_barras_modelos_1_datos.png'),
]

summary_rows = [
    ('1.0', 'FullGNN', '0.001058', 'MLP muy cerca; TinyGNN/PINN peor.'),
    ('0.5', 'FullGNN', '0.001062', 'Rendimiento casi igual que con 100%.'),
    ('0.2', 'FullGNN', '0.001052', 'No se observa degradación relevante.'),
    ('0.1', 'FullGNN', '0.001020', 'Mejor media del sweep.'),
    ('0.05', 'FullGNN', '0.001255', 'Ligera degradación, pero sigue estable.'),
    ('0.01', 'FullGNN', '0.001284', 'Mantiene buen resultado incluso con 1%.'),
]

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.6)
sec.bottom_margin = Inches(0.6)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Resumen del experimento: impacto de data_fraction')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(31, 78, 121)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Comparación de MLP, FullGNN, TinyGNN y TinyGNN+PINN en difusión de calor')
r.italic = True
r.font.size = Pt(11)

add_heading(doc, '1. Objetivo del experimento', 1)
doc.add_paragraph(
    'El objetivo fue analizar cómo afecta la reducción de datos de entrenamiento al rendimiento de los modelos. '
    'Para ello se varió el parámetro data_fraction y se evaluó el error de test, la violación física y el tiempo de inferencia.'
)

add_heading(doc, '2. Configuración resumida', 1)
for item in [
    'data_fraction probado: 1.0, 0.5, 0.2, 0.1, 0.05 y 0.01.',
    'Seeds utilizadas: 42, 43 y 44.',
    'Número de épocas: 50.',
    'Ejecución acelerada por GPU: CUDA.',
    'Modelos comparados: MLP, FullGNN, TinyGNN y TinyGNN + PINN.',
    'Métrica principal: test_mse medio. Menor valor indica mejor predicción.',
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading(doc, '3. Resultados principales', 1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['data_fraction', 'Mejor modelo', 'Test MSE medio', 'Comentario']):
    hdr[i].text = h
    set_cell_shading(hdr[i], 'D9EAF7')
    for p in hdr[i].paragraphs:
        for run in p.runs:
            run.bold = True
for row in summary_rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

add_heading(doc, '4. Gráficas', 1)
for caption, path in images:
    if not path.exists():
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_picture(str(path), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, '5. Conclusiones', 1)
conclusions = [
    'FullGNN fue el modelo más preciso y estable en todas las fracciones de datos evaluadas.',
    'MLP fue una baseline competitiva, con resultados cercanos a FullGNN, aunque normalmente algo peores.',
    'TinyGNN redujo drásticamente el número de parámetros, pero perdió bastante precisión frente a FullGNN y MLP.',
    'TinyGNN + PINN no mostró mejora clara respecto a TinyGNN en esta configuración; conviene revisar o ajustar physics_lambda y la formulación de la pérdida física.',
    'La reducción de datos no degradó tanto como se esperaba, probablemente porque el problema de difusión de calor es suave y regular.',
    'Para demostrar mejor la ventaja de enfoques físicos o robustos, sería recomendable probar escenarios más difíciles: más ruido, menos regularidad, distintas condiciones iniciales o mayor horizonte temporal.',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

add_heading(doc, '6. Conclusión global redactada', 1)
doc.add_paragraph(
    'Los resultados indican que FullGNN ofrece el mejor equilibrio entre precisión y robustez al reducir la cantidad de datos de entrenamiento. '
    'MLP se mantiene como una baseline fuerte, mientras que TinyGNN consigue una gran reducción de parámetros a costa de una pérdida notable de precisión. '
    'La variante TinyGNN+PINN no aporta una mejora significativa en esta configuración, por lo que su contribución debería estudiarse con un ajuste específico de la pérdida física o con escenarios más exigentes.'
)

doc.save(out_path)
print(out_path)
